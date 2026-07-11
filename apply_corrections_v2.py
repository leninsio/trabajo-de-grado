import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font_run(run, bold=False, italic=False, size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def format_body_paragraph(p, text, first_line_indent=0.5, left_indent=0.0):
    p.text = ""
    run = p.add_run(text)
    set_font_run(run)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Inches(first_line_indent)
    p.paragraph_format.left_indent = Inches(left_indent)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def format_quote_paragraph(p, text):
    format_body_paragraph(p, text, first_line_indent=0.0, left_indent=0.5)

def format_heading_paragraph(p, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p.text = ""
    run = p.add_run(text)
    set_font_run(run, bold=True)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.0)
    p.paragraph_format.left_indent = Inches(0.0)
    p.alignment = alignment

# =========================================================================
# I. APPLY CORRECTIONS TO MAIN THESIS
# =========================================================================
doc_path = "proyecto_de_grado/TRABAJO_LENIN_ALVARADO.docx"
doc = docx.Document(doc_path)

print("--- STARTING DOCX CORRECTIONS V2 ---")

# Helper to find paragraph indices by exact text
def find_p_index(text_query):
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == text_query:
            return idx
    return None

# 1. CHAPTER I: QUESTIONS & OBJECTIVES
print("Updating Chapter I Questions and Objectives...")
idx_q = find_p_index("Interrogantes de la Investigación")
if idx_q is not None:
    questions = [
        "¿Cuál es el estado actual de los procesos de registro, control de inventario y atención al cliente en la empresa Inversiones Más Por Tu Oro 17 C.A.?",
        "¿Qué requerimientos técnicos y funcionales son necesarios para la automatización y optimización de la gestión de inventario y control de préstamos prendarios de la empresa?",
        "¿Cómo debe diseñarse la arquitectura lógica del sistema, su base de datos relacional y las interfaces de usuario para garantizar la centralización y seguridad de la información?",
        "¿De qué manera el desarrollo e implementación del sistema de información propuesto optimizará el tiempo de respuesta y la confiabilidad de las operaciones de la empresa?"
    ]
    for offset, q_text in enumerate(questions):
        p_q = doc.paragraphs[idx_q + 1 + offset]
        format_body_paragraph(p_q, q_text, first_line_indent=0.5)

idx_o = find_p_index("Objetivos Específicos")
if idx_o is not None:
    objectives = [
        "Diagnosticar el estado actual de los procesos de registro, control de inventario y atención al cliente en la empresa Inversiones Más Por Tu Oro 17 C.A.",
        "Determinar los requerimientos técnicos y funcionales necesarios para la automatización y optimización de la gestión de inventario y control de préstamos prendarios.",
        "Diseñar la arquitectura lógica, la base de datos relacional y las interfaces de usuario del sistema de información para garantizar la centralización y seguridad de los datos.",
        "Desarrollar e implementar el sistema de información propuesto para optimizar el tiempo de respuesta y la confiabilidad de las operaciones de la empresa."
    ]
    for offset, o_text in enumerate(objectives):
        p_o = doc.paragraphs[idx_o + 1 + offset]
        format_body_paragraph(p_o, o_text, first_line_indent=0.5)


# 2. CHAPTER II: ANTECEDENTES (Continuous Narrative, Forward Order)
print("Updating Chapter II Antecedentes...")
idx_ant = find_p_index("Antecedentes de la Investigación")
if idx_ant is not None:
    idx_bt = find_p_index("Bases Teóricas")
    if idx_bt is not None:
        # Delete old schematic paragraphs
        for idx in range(idx_bt - 1, idx_ant, -1):
            p = doc.paragraphs[idx]
            p_element = p._p
            p_element.getparent().remove(p_element)
            
        # Re-fetch Bases Teóricas index
        idx_bt = find_p_index("Bases Teóricas")
        p_bt = doc.paragraphs[idx_bt]
        
        ant1_text = (
            "En relación con los antecedentes de la investigación, en primer lugar se cita el trabajo de Martínez, A. (2025), "
            "titulado \"Propuesta de un Sistema de Gestión de Inventario para Optimizar el Flujo de Metales Preciosos en Casas de Empeño\", "
            "presentado ante el Colegio Universitario de Administración y Mercadeo (CUAM), Extensión Caracas, para optar al título de "
            "Técnico Superior Universitario en Informática. El objetivo principal de dicho estudio fue proponer una solución tecnológica "
            "para mejorar el flujo físico y digital de metales preciosos dentro de una organización prendaria. Metodológicamente, se enmarcó "
            "bajo la modalidad de proyecto factible, con un diseño de campo de nivel descriptivo, aplicando una encuesta a una muestra "
            "de cinco (5) trabajadores. Entre sus conclusiones más relevantes, se determinó que la falta de centralización en el control "
            "de existencias eleva el riesgo de pérdidas financieras. Este trabajo aporta a la presente investigación una base metodológica "
            "sólida para el levantamiento de requerimientos funcionales y un modelo entidad-relación que sirve de referencia directa para "
            "estructurar las tablas de inventario en la base de datos relacional."
        )
        ant2_text = (
            "En segundo lugar, se analizó la investigación de Colina, J. (2024), denominada \"Análisis y Desarrollo de Módulos de Cálculo "
            "para la Auditoría de Créditos Prendarios y Mitigación de Riesgos Financieros\", desarrollada en la Universidad Alejandro de "
            "Humboldt (UAH) para optar al título de Ingeniero en Informática. La investigación tuvo como propósito diseñar algoritmos seguros "
            "para el cálculo automático de intereses devengados por contratos de empeño. Se trató de una investigación tecnológica y descriptiva "
            "con apoyo de campo. La conclusión fundamental de la investigación demostró que la automatización de fórmulas reduce a cero el "
            "margen de error humano en operaciones matemáticas complejas de amortización. La relación con el presente proyecto reside en que "
            "proporciona los fundamentos algorítmicos y matemáticos para estructurar los módulos de cálculo de tasas de interés y fechas de "
            "vencimiento del préstamo prendario de oro y plata."
        )
        ant3_text = (
            "Por último, se consideró el trabajo de Naranjo, E. (2023), titulado \"Impacto de la Migración de Datos de Hojas de Cálculo a "
            "un Ambiente Web Centralizado en la Productividad Administrativa\", presentado ante el Instituto Universitario de Tecnología de "
            "Administración Industrial (IUTA). El objetivo general fue evaluar las mejoras en la eficiencia operativa tras migrar registros "
            "tradicionales de Excel a una plataforma centralizada en la web. La metodología correspondió a un diseño transeccional descriptivo "
            "apoyado en una muestra de ocho (8) usuarios. La investigación concluyó que los entornos web centralizados reducen el tiempo de "
            "respuesta administrativa en más de un 60% y aumentan la consistencia de los datos. Este antecedente constituye una justificación "
            "metodológica clave para nuestro estudio, al validar técnicamente el reemplazo de los libros de cálculo manuales descentralizados "
            "por el software interactivo propuesto."
        )
        
        # Insert in FORWARD order
        p_ant1 = p_bt.insert_paragraph_before()
        format_body_paragraph(p_ant1, ant1_text)
        
        p_ant2 = p_bt.insert_paragraph_before()
        format_body_paragraph(p_ant2, ant2_text)
        
        p_ant3 = p_bt.insert_paragraph_before()
        format_body_paragraph(p_ant3, ant3_text)


# 3. CHAPTER II: BASES TEÓRICAS (Extended with Subsections, Forward Order)
print("Updating Chapter II Bases Teóricas...")
idx_bt = find_p_index("Bases Teóricas")
if idx_bt is not None:
    # Delete the single paragraph following Bases Teóricas
    p_to_del = doc.paragraphs[idx_bt + 1]
    p_to_del_element = p_to_del._p
    p_to_del_element.getparent().remove(p_to_del_element)
    
    # Re-fetch Bases Legales index
    idx_bl = find_p_index("Bases Legales")
    p_bl = doc.paragraphs[idx_bl]
    
    sub1_p1 = (
        "Los sistemas de información representan estructuras integradas por recursos humanos, tecnológicos y normativos, "
        "diseñados con el propósito de recopilar, almacenar y transformar datos en información útil que sirva de apoyo en "
        "la toma de decisiones operativas y estratégicas de las organizaciones. En el contexto de las microempresas, "
        "su implantación marca el paso desde procesos manuales o de ofimática descentralizada hacia flujos de trabajo "
        "centralizados y controlados."
    )
    sub1_quote = (
        "Un sistema de información se define como un conjunto de componentes interrelacionados que recolectan, "
        "procesan, almacenan y distribuyen información para apoyar la toma de decisiones y el control en una organización. "
        "Los sistemas de información ayudan a los gerentes y trabajadores a analizar problemas, visualizar asuntos "
        "complejos y crear nuevos productos. (Kendall y Kendall, 2011, p. 12)"
    )
    sub1_p2 = (
        "De este modo, para la empresa Inversiones Más Por Tu Oro 17 C.A., el sistema propuesto no constituye únicamente "
        "un repositorio de datos, sino un núcleo operativo que articula la gestión del inventario y la valoración de "
        "prendas en tiempo real, agilizando el flujo diario y eliminando la duplicidad en el procesamiento."
    )
    
    sub2_p1 = (
        "Una base de datos es una colección organizada y estructurada de información persistente, gestionada mediante un "
        "software especializado que garantiza el acceso rápido, seguro y coherente de los usuarios. En el desarrollo "
        "de software, la adopción del modelo relacional es fundamental para asegurar que los datos no sufran de "
        "inconsistencias o sobreescrituras accidentales."
    )
    sub2_quote = (
        "Un sistema de base de datos relacional es un sistema en el cual los datos se perciben por los usuarios como "
        "tablas lógicas, y solo como tablas relacionales, permitiendo realizar consultas complejas mediante la correspondencia "
        "de atributos comunes entre diferentes entidades lógicas. (Date, 2004, p. 25)"
    )
    sub2_p2 = (
        "Para el control de los préstamos prendarios, el modelo relacional permite vincular la tabla de clientes con la tabla "
        "de contratos de empeño y la tabla de inventario de metales. Esto asegura la integridad referencial de los registros: "
        "no puede existir una transacción de empeño sin un deudor registrado ni un metal precioso sin su correspondiente "
        "ficha técnica de kilataje y peso en gramos, lo cual es imposible de lograr con hojas de cálculo aisladas de Excel."
    )
    
    sub3_p1 = (
        "La ingeniería de software comprende el conjunto de metodologías, técnicas y herramientas disciplinadas destinadas al "
        "desarrollo e implantación de aplicaciones informáticas eficientes y escalables. Durante el ciclo de vida del software, "
        "el modelado de la interfaz y la definición de requerimientos funcionales constituyen los pilares del diseño lógico."
    )
    sub3_quote = (
        "El modelo de construcción de prototipos ayuda al desarrollador y al cliente a entender mejor los requisitos del sistema "
        "cuando los objetivos generales del software están definidos, pero los detalles técnicos de entrada y salida aún no "
        "se comprenden en su totalidad por parte de los usuarios finales. (Pressman, 2010, p. 45)"
    )
    sub3_p2 = (
        "Conforme a esto, este proyecto implementa un prototipo de interfaz interactiva desplegado en web, lo cual permite al "
        "personal operativo de la empresa validar el diseño de las pantallas y flujos de consulta antes de consolidar el "
        "backend. Esto minimiza la resistencia al cambio tecnológico y garantiza que el software se ajuste a las necesidades "
        "reales del puesto de trabajo."
    )
    
    sub4_p1 = (
        "La gestión de inventario involucra el control sistemático de los bienes muebles que ingresan y salen de una organización. "
        "En el caso particular de metales preciosos (oro y plata), la custodia exige una rigurosidad extrema debido a la alta valorización "
        "de los activos y a las fluctuaciones constantes de su precio en el mercado financiero."
    )
    sub4_p2 = (
        "Por otro lado, el préstamo prendario representa una operación contractual donde el deudor entrega un bien tangible "
        "(la prenda de oro o plata) al acreedor como garantía real del crédito recibido. La administración de este proceso demanda "
        "cálculos precisos de tasas de interés y un control estricto de las fechas de vencimiento para evitar demoras operativas o "
        "pérdidas por custodia deficiente. El sistema de información automatiza estas variables financieras para dar transparencia "
        "a ambas partes de la relación contractual."
    )
    
    # Insert in FORWARD order (so they appear: Sub 1, Sub 2, Sub 3, Sub 4)
    # Sub 1
    p_s1_h = p_bl.insert_paragraph_before()
    format_heading_paragraph(p_s1_h, "Sistemas de Información")
    p_s1_p1 = p_bl.insert_paragraph_before()
    format_body_paragraph(p_s1_p1, sub1_p1)
    p_s1_q = p_bl.insert_paragraph_before()
    format_quote_paragraph(p_s1_q, sub1_quote)
    p_s1_p2 = p_bl.insert_paragraph_before()
    format_body_paragraph(p_s1_p2, sub1_p2)
    
    # Sub 2
    p_s2_h = p_bl.insert_paragraph_before()
    format_heading_paragraph(p_s2_h, "Bases de Datos Relacionales y Centralización")
    p_s2_p1 = p_bl.insert_paragraph_before()
    format_body_paragraph(p_s2_p1, sub2_p1)
    p_s2_q = p_bl.insert_paragraph_before()
    format_quote_paragraph(p_s2_q, sub2_quote)
    p_s2_p2 = p_bl.insert_paragraph_before()
    format_body_paragraph(p_s2_p2, sub2_p2)
    
    # Sub 3
    p_s3_h = p_bl.insert_paragraph_before()
    format_heading_paragraph(p_s3_h, "Ingeniería de Software y Modelado de Interfaces")
    p_s3_p1 = p_bl.insert_paragraph_before()
    format_body_paragraph(p_s3_p1, sub3_p1)
    p_s3_q = p_bl.insert_paragraph_before()
    format_quote_paragraph(p_s3_q, sub3_quote)
    p_s3_p2 = p_bl.insert_paragraph_before()
    format_body_paragraph(p_s3_p2, sub3_p2)
    
    # Sub 4
    p_s4_h = p_bl.insert_paragraph_before()
    format_heading_paragraph(p_s4_h, "Gestión de Inventario de Metales Preciosos y Préstamos Prendarios")
    p_s4_p1 = p_bl.insert_paragraph_before()
    format_body_paragraph(p_s4_p1, sub4_p1)
    p_s4_p2 = p_bl.insert_paragraph_before()
    format_body_paragraph(p_s4_p2, sub4_p2)


# 4. CHAPTER II: BASES LEGALES (Extended Relational Analysis, Ley Orgánica de Minas 2026, Forward Order)
print("Updating Chapter II Bases Legales...")
idx_bl = find_p_index("Bases Legales")
if idx_bl is not None:
    # Delete the single paragraph following Bases Legales
    p_to_del = doc.paragraphs[idx_bl + 1]
    p_to_del_element = p_to_del._p
    p_to_del_element.getparent().remove(p_to_del_element)
    
    # Re-fetch Cuadro de Operacionalización de Variables index
    idx_cov = find_p_index("Cuadro de Operacionalización de Variables")
    p_cov = doc.paragraphs[idx_cov]
    
    bl_intro = (
        "El sustento jurídico de esta investigación se enmarca en las leyes y normativas venezolanas que regulan el uso "
        "de tecnologías de información y controlan la actividad comercial y financiera de custodia y préstamo sobre bienes "
        "muebles. A continuación se detallan las bases legales y el análisis relacional que las vincula directamente con el sistema propuesto."
    )
    
    bl_c1 = (
        "Constitución de la República Bolivariana de Venezuela (1999):\n"
        "El Artículo 110 establece que el Estado reconocerá el interés público de la ciencia, la tecnología, el conocimiento, "
        "la información y sus servicios asociados como instrumentos fundamentales para el desarrollo económico, social y político del país. "
        "Asimismo, indica que el Estado garantizará los recursos necesarios para el fomento de estas actividades."
    )
    bl_c1_analysis = (
        "Análisis de relación: El desarrollo de este software de gestión representa una materialización a microescala de los objetivos "
        "del Artículo 110, al incorporar innovación tecnológica en un sector comercial tradicional (comercio prendario) para "
        "modernizar sus procesos administrativos. El proyecto demuestra cómo el uso de herramientas tecnológicas libres fomenta "
        "el desarrollo del comercio y mejora la eficiencia de los servicios locales, alineándose con las políticas públicas de fomento tecnológico."
    )
    
    bl_c2 = (
        "Ley Especial Contra los Delitos Informáticos (2001):\n"
        "Esta ley tiene por objeto la protección integral de los sistemas que utilicen tecnologías de información, así como la prevención "
        "y sanción de los delitos cometidos contra o mediante el uso de tales sistemas. Artículos clave regulan el acceso indebido (Art. 6) "
        "y la manipulación de sistemas protegidos o alteración de datos (Art. 7)."
    )
    bl_c2_analysis = (
        "Análisis de relación: El sistema propuesto implementa controles de seguridad física y lógica para cumplir con esta normativa. "
        "Se diseñan módulos de acceso restringido mediante perfiles de usuario y contraseñas cifradas, limitando las acciones operativas "
        "(Gerente, Administrador, Cajero) y resguardando la base de datos contra accesos no autorizados. Adicionalmente, el diseño contempla "
        "respaldos automáticos de datos para evitar la pérdida o alteración ilícita de los registros de contratos de empeño y garantías reales."
    )
    
    bl_c3 = (
        "Código de Comercio de Venezuela (1955):\n"
        "El Artículo 32 establece la obligación de todo comerciante de llevar la contabilidad mercantil de forma organizada, debiendo llevar "
        "el Libro Diario, el Libro Mayor y el de Inventarios, en idioma castellano y de forma cronológica."
    )
    bl_c3_analysis = (
        "Análisis de relación: La aplicación web para la gestión de inventario y préstamos prendarios actúa como un soporte técnico fundamental "
        "para el cumplimiento de las obligaciones del Código de Comercio. Al centralizar los registros en una base de datos relacional, el sistema "
        "genera de forma automática reportes cronológicos de transacciones de compra, custodia y préstamos activos, garantizando la fidelidad y "
        "el orden del inventario sin los riesgos de inconsistencia aritmética que caracterizan a los registros manuales en hojas de cálculo."
    )
    
    bl_c4 = (
        "Ley Orgánica de Minas (Gaceta Oficial N° 6.802 Extraordinario del 15 de abril de 2026):\n"
        "Esta ley regula de forma consolidada las actividades mineras y sus conexas, incluyendo el procesamiento, tenencia, "
        "transporte y comercialización de metales preciosos dentro del territorio nacional, estableciendo además el derecho "
        "preferente de compra por parte del Banco Central de Venezuela (BCV). Asimismo, se complementa con las normativas "
        "sobre prevención de legitimación de capitales bajo la Ley Orgánica contra la Delincuencia Organizada y Financiamiento al Terrorismo (2012)."
    )
    bl_c4_analysis = (
        "Análisis de relación: El sistema propuesto se alinea directamente con las exigencias de fiscalización y registro de la Ley Orgánica de "
        "Minas y las normativas de prevención de legitimación de capitales. Al automatizar la ficha de registro de los contratos prendarios, "
        "el software almacena y resguarda de forma segura los datos exigidos para la trazabilidad de los metales adquiridos o recibidos en garantía: "
        "peso exacto en gramos, kilataje, descripción detallada del bien y la identificación plenamente verificada del deudor (C.I., firma y "
        "huella digital). De esta manera, el sistema asegura que la empresa Inversiones Más Por Tu Oro 17 C.A. mantenga registros fidedignos "
        "y auditables que demuestren que los metales custodiados corresponden a piezas de uso personal y comercialización lícita, "
        "facilitando la transparencia operativa ante cualquier fiscalización de los entes gubernamentales."
    )
    
    # Insert in FORWARD order
    p_intro = p_cov.insert_paragraph_before()
    format_body_paragraph(p_intro, bl_intro)
    
    p_l1 = p_cov.insert_paragraph_before()
    format_body_paragraph(p_l1, bl_c1)
    p_l1_a = p_cov.insert_paragraph_before()
    format_body_paragraph(p_l1_a, bl_c1_analysis)
    
    p_l2 = p_cov.insert_paragraph_before()
    format_body_paragraph(p_l2, bl_c2)
    p_l2_a = p_cov.insert_paragraph_before()
    format_body_paragraph(p_l2_a, bl_c2_analysis)
    
    p_l3 = p_cov.insert_paragraph_before()
    format_body_paragraph(p_l3, bl_c3)
    p_l3_a = p_cov.insert_paragraph_before()
    format_body_paragraph(p_l3_a, bl_c3_analysis)
    
    p_l4 = p_cov.insert_paragraph_before()
    format_body_paragraph(p_l4, bl_c4)
    p_l4_a = p_cov.insert_paragraph_before()
    format_body_paragraph(p_l4_a, bl_c4_analysis)


# 5. CHAPTER III: TIPO Y DISEÑO DE LA INVESTIGACIÓN (Delimited real, Forward Order)
print("Updating Chapter III Tipo y Diseño...")
idx_met = find_p_index("Tipo y Diseño de la Investigación")
if idx_met is not None:
    idx_pm = find_p_index("Población y Muestra")
    if idx_pm is not None:
        # Delete old paragraphs
        for idx in range(idx_pm - 1, idx_met, -1):
            p = doc.paragraphs[idx]
            p_element = p._p
            p_element.getparent().remove(p_element)
            
        # Re-fetch Population index
        idx_pm = find_p_index("Población y Muestra")
        p_pm = doc.paragraphs[idx_pm]
        
        m_intro = (
            "De acuerdo con la naturaleza del problema planteado y los objetivos propuestos para la resolución del control de inventario "
            "y registro de contratos prendarios en la empresa Inversiones Más Por Tu Oro 17 C.A., el estudio se clasifica metodológicamente "
            "bajo la modalidad de Proyecto Factible, apoyado en un diseño de Campo de nivel Descriptivo. A continuación se detallan las especificaciones técnicas:"
        )
        m_factible = (
            "Modalidad de la Investigación (Proyecto Factible):\n"
            "Este estudio consiste en la formulación de una propuesta tecnológica viable para resolver la inoperancia administrativa detectada en "
            "el flujo de inventario de la empresa. Al respecto, el manual de la Universidad Pedagógica Experimental Libertador (UPEL, 2016) define "
            "el Proyecto Factible como la investigación, elaboración y desarrollo de una propuesta de un modelo operativo viable para solucionar "
            "problemas, requerimientos o necesidades de organizaciones o grupos sociales (p. 21). El sistema de información web para la gestión de "
            "inventario y control de préstamos representa un modelo tecnológico y operativo funcional que atiende directamente una necesidad práctica e institucional."
        )
        m_diseno = (
            "Diseño de la Investigación (Campo):\n"
            "El diseño de la investigación es de Campo, dado que los datos sobre la situación operativa, los tiempos de atención y los requerimientos del "
            "sistema se recolectan directamente en el entorno físico de la empresa Inversiones Más Por Tu Oro 17 C.A., en su contexto cotidiano, sin "
            "manipular ni controlar intencionalmente ninguna de las variables operativas del estudio. Según Arias (2012), la investigación de campo consiste "
            "en la recolección de datos directamente de los sujetos investigados, o de la realidad donde ocurren los hechos (sin manipular o controlar variable "
            "alguna) (p. 31). Este diseño es el técnicamente adecuado porque permite contrastar la teoría de la base de datos con los requerimientos reales recopilados del personal."
        )
        m_nivel = (
            "Nivel de la Investigación (Descriptivo):\n"
            "El estudio posee un nivel Descriptivo, en vista de que se orienta a caracterizar, analizar e identificar de forma rigurosa las características del "
            "sistema de registro manual actual, sus fallas, retrasos y deficiencias en el cálculo aritmético, para posteriormente definir con total precisión técnica "
            "los requerimientos que el nuevo software debe poseer. Al respecto, Arias (2012) señala que la investigación descriptiva consiste en la caracterización "
            "de un hecho, fenómeno, individuo o grupo, con el fin de establecer su estructura o comportamiento (p. 24). Se descarta explícitamente el nivel explicativo, "
            "ya que el presente proyecto no persigue demostrar hipótesis correlacionales de causa y efecto mediante experimentos controlados, sino proponer y desarrollar "
            "una solución de ingeniería de software viable ante una problemática delimitada."
        )
        
        # Insert in FORWARD order
        p_m_intro = p_pm.insert_paragraph_before()
        format_body_paragraph(p_m_intro, m_intro)
        
        p_m_factible = p_pm.insert_paragraph_before()
        format_body_paragraph(p_m_factible, m_factible)
        
        p_m_diseno = p_pm.insert_paragraph_before()
        format_body_paragraph(p_m_diseno, m_diseno)
        
        p_m_nivel = p_pm.insert_paragraph_before()
        format_body_paragraph(p_m_nivel, m_nivel)


# 6. SIGNATURE LINES (Add __________________________________ above Br. Lenin Alvarado)
print("Adding signature lines in main thesis...")
# We must scan paragraphs for "Br. Lenin Alvarado"
# Since we might add paragraphs, we collect indices and iterate in reverse
sig_indices = []
for idx, p in enumerate(doc.paragraphs):
    if "Br. Lenin Alvarado" in p.text:
        sig_indices.append(idx)

for sig_idx in reversed(sig_indices):
    p_sig = doc.paragraphs[sig_idx]
    # Check if the paragraph before or two paragraphs before has the line
    p_before1 = doc.paragraphs[sig_idx - 1]
    p_before2 = doc.paragraphs[sig_idx - 2]
    if "________" not in p_before1.text and "________" not in p_before2.text:
        p_line = p_sig.insert_paragraph_before()
        p_line.text = "__________________________________"
        p_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_line.paragraph_format.space_before = Pt(12)
        p_line.paragraph_format.space_after = Pt(2)
        p_line.paragraph_format.first_line_indent = Inches(0.0)
        p_line.paragraph_format.left_indent = Inches(0.0)
        # Apply Times New Roman
        for r in p_line.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)

# 7. SYNC VARIABLES TABLES IN MAIN THESIS
print("Syncing Variables Tables in Main Thesis...")
new_table_objectives = [
    "Diagnosticar el estado de registro, control de inventario y atención al cliente.",
    "Determinar los requerimientos técnicos y funcionales del sistema propuesto.",
    "Diseñar la arquitectura lógica, base de datos relacional e interfaces.",
    "Desarrollar e implementar el sistema de información propuesto."
]

def sync_doc_tables(doc_obj):
    updated_count = 0
    for idx, table in enumerate(doc_obj.tables):
        found = False
        for row in table.rows:
            if "Diagnosticar" in row.cells[0].text:
                found = True
                break
        if found:
            print(f"  Updating table {idx}...")
            row_idx = 1
            for o_txt in new_table_objectives:
                if row_idx < len(table.rows):
                    cell = table.rows[row_idx].cells[0]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(o_txt)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    row_idx += 1
            updated_count += 1
    return updated_count

updated_tables = sync_doc_tables(doc)
print(f"Updated {updated_tables} tables in main thesis.")
doc.save(doc_path)
print("Saved corrected TRABAJO_LENIN_ALVARADO.docx.")


# =========================================================================
# II. APPLY CORRECTIONS TO VALIDATION KIT
# =========================================================================
print("\nUpdating Instrument_y_Validacion_Expertos.docx...")
kit_path = "proyecto_de_grado/Instrumento_y_Validacion_Expertos.docx"
kit_doc = docx.Document(kit_path)

# Update variables table in kit
updated_kit_tables = sync_doc_tables(kit_doc)
print(f"Updated {updated_kit_tables} tables in validation kit.")
kit_doc.save(kit_path)
print("Saved corrected Instrumento_y_Validacion_Expertos.docx.")


# =========================================================================
# III. SYNC CUESTIONARIO EMPLEADOS (Socio-demographics to match thesis exactly)
# =========================================================================
print("\nUpdating Cuestionario_Empleados.docx...")
emp_path = "proyecto_de_grado/Cuestionario_Empleados.docx"
emp_doc = docx.Document(emp_path)

# Locate paragraph containing "DATOS DEL TRABAJADOR" or similar
target_idx = None
for idx, p in enumerate(emp_doc.paragraphs):
    if "DATOS DEL TRABAJADOR" in p.text.upper():
        target_idx = idx
        break

if target_idx is not None:
    print(f"Found workers data section at index {target_idx}")
    # Delete this paragraph and the following two paragraphs (which are the blanks)
    # P3: DATOS DEL TRABAJADOR
    # P4: Cargo que desempeña
    # P5: Años de experiencia
    # We delete indices: target_idx + 2, target_idx + 1, target_idx
    for idx in range(target_idx + 2, target_idx - 1, -1):
        if idx < len(emp_doc.paragraphs):
            p = emp_doc.paragraphs[idx]
            p_element = p._p
            p_element.getparent().remove(p_element)
            
    # Now insert the exact checkbox format at the target_idx position
    # Let's locate the next element in the doc body (Table 1 which is the questions table)
    # We will insert paragraphs before Table 1:
    table_q = emp_doc.tables[1]
    
    # We insert the paragraphs in forward order
    p_title = table_q.rows[0].cells[0].paragraphs[0].insert_paragraph_before()
    p_title.text = "I. DATOS SOCIODEMOGRÁFICOS"
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.line_spacing = 1.5
    for r in p_title.runs:
         set_font_run(r, bold=True, size=11)
         
    p_cargo = table_q.rows[0].cells[0].paragraphs[0].insert_paragraph_before()
    p_cargo.text = "Cargo que ocupa: [ ] Gerencia   [ ] Administrativo   [ ] Operativo/Caja"
    p_cargo.paragraph_format.space_before = Pt(0)
    p_cargo.paragraph_format.space_after = Pt(6)
    p_cargo.paragraph_format.line_spacing = 1.5
    for r in p_cargo.runs:
         set_font_run(r, size=11)
         
    p_exp = table_q.rows[0].cells[0].paragraphs[0].insert_paragraph_before()
    p_exp.text = "Años de experiencia: [ ] Menos de 1 año   [ ] 1 a 3 años   [ ] Más de 3 años"
    p_exp.paragraph_format.space_before = Pt(0)
    p_exp.paragraph_format.space_after = Pt(6)
    p_exp.paragraph_format.line_spacing = 1.5
    for r in p_exp.runs:
         set_font_run(r, size=11)
         
    p_edu = table_q.rows[0].cells[0].paragraphs[0].insert_paragraph_before()
    p_edu.text = "Nivel Educativo: [ ] Bachiller   [ ] TSU   [ ] Universitario"
    p_edu.paragraph_format.space_before = Pt(0)
    p_edu.paragraph_format.space_after = Pt(12)
    p_edu.paragraph_format.line_spacing = 1.5
    for r in p_edu.runs:
         set_font_run(r, size=11)
         
    p_title2 = table_q.rows[0].cells[0].paragraphs[0].insert_paragraph_before()
    p_title2.text = "II. CUESTIONARIO"
    p_title2.paragraph_format.space_before = Pt(6)
    p_title2.paragraph_format.space_after = Pt(6)
    p_title2.paragraph_format.line_spacing = 1.5
    for r in p_title2.runs:
         set_font_run(r, bold=True, size=11)

emp_doc.save(emp_path)
print("Saved corrected Cuestionario_Empleados.docx.")

print("\n--- ALL CORRECTIONS COMPLETED V2 ---")
