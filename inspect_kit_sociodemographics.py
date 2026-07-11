import docx

kit = docx.Document("proyecto_de_grado/Instrumento_y_Validacion_Expertos.docx")
print("--- ALL PARAGRAPHS IN VALIDATION KIT ---")
for idx, p in enumerate(kit.paragraphs):
    if "SOCIODEMOGRÁFICOS" in p.text.upper() or "DATOS DEL TRABAJADOR" in p.text.upper() or "SOCIODEMO" in p.text.upper():
        print(f"Found in paragraph {idx}: '{p.text}'")
        for offset in range(1, 8):
            c_idx = idx + offset
            if c_idx < len(kit.paragraphs):
                print(f"  P{c_idx}: '{kit.paragraphs[c_idx].text.strip()}'")
        break
