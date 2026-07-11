import docx

doc = docx.Document("proyecto_de_grado/TRABAJO_LENIN_ALVARADO.docx")
print("--- CHAPTER I OBJECTIVES AND QUESTIONS ---")

# Let's search for headings like "Objetivo General", "Objetivos Específicos", "Interrogantes", "Planteamiento"
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Interrogante" in text or "Objetivo General" in text or "Objetivos Específicos" in text or "Objetivo de la" in text:
        print(f"P{idx}: '{p.text}'")
        # Print surrounding paragraphs
        for offset in range(1, 8):
            c_idx = idx + offset
            if c_idx < len(doc.paragraphs):
                print(f"  P{c_idx}: '{doc.paragraphs[c_idx].text.strip()}'")
