import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font_run(run, bold=False, size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def format_body_paragraph(p, text, first_line_indent=0.5, bold=False):
    p.text = ""
    run = p.add_run(text)
    set_font_run(run, bold=bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(first_line_indent)
    p.paragraph_format.left_indent = Inches(0.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc_path = "proyecto_de_grado/TRABAJO_LENIN_ALVARADO.docx"
doc = docx.Document(doc_path)

print("--- FIXING ANTECEDENTES PARAGRAPH ORDER ---")

# Let's locate the paragraphs by index first to be sure
idx_martinez = None
idx_naranjo = None
idx_colina = None

for idx, p in enumerate(doc.paragraphs):
    if "Martínez, A. (2025)" in p.text:
        idx_martinez = idx
    elif "Colina, J. (2024)" in p.text:
        idx_colina = idx
    elif "Naranjo, E. (2023)" in p.text:
        idx_naranjo = idx

if idx_martinez is not None and idx_colina is not None and idx_naranjo is not None:
    print(f"Current indices: Naranjo={idx_naranjo}, Colina={idx_colina}, Martínez={idx_martinez}")
    
    # Store text
    text_martinez = doc.paragraphs[idx_martinez].text
    text_colina = doc.paragraphs[idx_colina].text
    text_naranjo = doc.paragraphs[idx_naranjo].text
    
    # Rearrange them in chronological/logical order: Martínez first (P181), Colina second (P182), Naranjo third (P183)
    # The current indices are: P181 (Naranjo), P182 (Colina), P183 (Martínez)
    # We want:
    # P181 -> Martínez (2025)
    # P182 -> Colina (2024)
    # P183 -> Naranjo (2023)
    
    format_body_paragraph(doc.paragraphs[181], text_martinez, first_line_indent=0.5)
    format_body_paragraph(doc.paragraphs[182], text_colina, first_line_indent=0.5)
    format_body_paragraph(doc.paragraphs[183], text_naranjo, first_line_indent=0.5)

doc.save(doc_path)
print("Successfully corrected Antecedentes paragraph order in main thesis.")
