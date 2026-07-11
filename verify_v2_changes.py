import docx

doc = docx.Document("proyecto_de_grado/TRABAJO_LENIN_ALVARADO.docx")
emp = docx.Document("proyecto_de_grado/Cuestionario_Empleados.docx")

print("=== VERIFYING BASES TEORICAS IN THESIS ===")
idx_bt = None
for idx, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Bases Teóricas":
        idx_bt = idx
        break

if idx_bt is not None:
    for offset in range(1, 15):
        c_idx = idx_bt + offset
        if c_idx < len(doc.paragraphs):
            print(f"P{c_idx}: '{doc.paragraphs[c_idx].text[:80].strip()}...'")

print("\n=== VERIFYING BASES LEGALES IN THESIS ===")
idx_bl = None
for idx, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Bases Legales":
        idx_bl = idx
        break

if idx_bl is not None:
    for offset in range(1, 12):
        c_idx = idx_bl + offset
        if c_idx < len(doc.paragraphs):
            print(f"P{c_idx}: '{doc.paragraphs[c_idx].text[:80].strip()}...'")

print("\n=== VERIFYING SIGNATURES IN THESIS ===")
for idx, p in enumerate(doc.paragraphs):
    if "Br. Lenin Alvarado" in p.text:
        print(f"Found signature at P{idx}: '{p.text}'")
        print(f"  P{idx-1}: '{doc.paragraphs[idx-1].text}'")

print("\n=== VERIFYING CUESTIONARIO EMPLEADOS SOCIODEMOGRAPHICS ===")
for idx, p in enumerate(emp.paragraphs):
    if "SOCIODEMOGRÁFICOS" in p.text or "Cargo que ocupa" in p.text or "Nivel Educativo" in p.text:
        print(f"P{idx}: '{p.text}'")
