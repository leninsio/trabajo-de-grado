import docx

def find_lenin_signatures(file_path):
    print(f"\n--- Checking signatures in {file_path} ---")
    doc = docx.Document(file_path)
    for idx, p in enumerate(doc.paragraphs):
        if "Br. Lenin Alvarado" in p.text:
            print(f"Found 'Br. Lenin Alvarado' at index {idx}:")
            # Print 3 paragraphs before and after
            start = max(0, idx - 3)
            end = min(len(doc.paragraphs), idx + 3)
            for i in range(start, end):
                print(f"  P{i}: '{doc.paragraphs[i].text.strip()}'")

find_lenin_signatures("proyecto_de_grado/TRABAJO_LENIN_ALVARADO.docx")
find_lenin_signatures("proyecto_de_grado/Instrumento_y_Validacion_Expertos.docx")
find_lenin_signatures("proyecto_de_grado/Cuestionario_Empleados.docx")
