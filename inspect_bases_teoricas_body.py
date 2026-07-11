import docx

doc = docx.Document("proyecto_de_grado/TRABAJO_LENIN_ALVARADO.docx")
print("--- ALL BASES TEORICAS OCCURRENCES ---")
for idx, p in enumerate(doc.paragraphs):
    if "Bases Teóricas" in p.text and idx > 120:
        print(f"Found heading at P{idx}: '{p.text}'")
        # Print subsequent paragraphs to read the text
        for offset in range(1, 30):
            c_idx = idx + offset
            if c_idx < len(doc.paragraphs):
                print(f"  P{c_idx}: '{doc.paragraphs[c_idx].text.strip()}'")
        break
