import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font_run(run, bold=False, italic=False, size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def format_body_paragraph(p, text, first_line_indent=0.5, left_indent=0.0):
    p.text = ""
    run = p.add_run(text)
    set_font_run(run)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Inches(first_line_indent)
    p.paragraph_format.left_indent = Inches(left_indent)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc_path = "proyecto_de_grado/TRABAJO_LENIN_ALVARADO.docx"
doc = docx.Document(doc_path)

print("--- UPDATING LEGAL BASES TO 2026 REGULATION ---")

found_law = False
found_analysis = False

# We locate and replace the text of the 4th legal base and its analysis
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Reserva al Estado las Actividades de Exploración y Explotación del Oro" in text or "Decreto con Rango, Valor y Fuerza de Ley que Reserva" in text:
        print(f"Found old 4th law paragraph at index {idx}")
        new_law_text = (
            "Ley Orgánica de Minas (Gaceta Oficial N° 6.802 Extraordinario del 15 de abril de 2026):\n"
            "Esta ley regula de forma consolidada las actividades mineras y sus conexas, incluyendo el procesamiento, tenencia, "
            "transporte y comercialización de metales preciosos dentro del territorio nacional, estableciendo además el derecho "
            "preferente de compra por parte del Banco Central de Venezuela (BCV). Asimismo, se complementa con las normativas "
            "sobre prevención de legitimación de capitales bajo la Ley Orgánica contra la Delincuencia Organizada y Financiamiento al Terrorismo (2012)."
        )
        format_body_paragraph(p, new_law_text)
        found_law = True
        
    elif "específicamente parametrizado para almacenar las variables requeridas por este marco legal" in text or "Inversiones Más Por Tu Oro 17 C.A. mantenga registros fidedignos" in text or "De este modo, la empresa Inversiones Más Por Tu Oro" in text and "ley de metales preciosos" in text:
        print(f"Found old 4th analysis paragraph at index {idx}")
        new_analysis_text = (
            "Análisis de relación: El sistema propuesto se alinea directamente con las exigencias de fiscalización y registro de la Ley Orgánica de "
            "Minas y las normativas de prevención de legitimación de capitales. Al automatizar la ficha de registro de los contratos prendarios, "
            "el software almacena y resguarda de forma segura los datos exigidos para la trazabilidad de los metales adquiridos o recibidos en garantía: "
            "peso exacto en gramos, kilataje, descripción detallada del bien y la identificación plenamente verificada del deudor (C.I., firma y "
            "huella digital). De esta manera, el sistema asegura que la empresa Inversiones Más Por Tu Oro 17 C.A. mantenga registros fidedignos "
            "y auditables que demuestren que los metales custodiados corresponden a piezas de uso personal y comercialización lícita, "
            "facilitando la transparencia operativa ante cualquier fiscalización de los entes gubernamentales."
        )
        format_body_paragraph(p, new_analysis_text)
        found_analysis = True

if found_law and found_analysis:
    doc.save(doc_path)
    print("Successfully updated bases legales with current 2026 regulations.")
else:
    print("Error: Could not find target paragraphs for 4th law and analysis!")
    # Let's search by keywords
    print("\nAlternative keyword search:")
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if "Preciosos" in txt and idx > 190 and idx < 220:
            print(f"  P{idx}: '{txt[:80]}...'")
