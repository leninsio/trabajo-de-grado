import docx

doc = docx.Document("proyecto_de_grado/TRABAJO_LENIN_ALVARADO.docx")
print("--- ANTECEDENTES SECTION INSPECT ---")
found = False
for idx, p in enumerate(doc.paragraphs):
    if "Antecedentes de la" in p.text or "ANTECEDENTES DE LA" in p.text:
        found = True
        print(f"Found heading at P{idx}: '{p.text}'")
        # Print subsequent paragraphs to read the text
        for offset in range(1, 35):
            c_idx = idx + offset
            if c_idx < len(doc.paragraphs):
                print(f"  P{c_idx}: '{doc.paragraphs[c_idx].text.strip()}'")
        break
if not found:
    print("Could not find Antecedentes section!")
